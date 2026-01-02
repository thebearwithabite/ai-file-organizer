#!/usr/bin/env python3
"""
Content Extraction Pipeline for Local LLM Librarian
Extracts text content from various file formats for indexing and search
"""

import os
import json
import sqlite3
from pathlib import Path
from typing import Dict, List, Optional, Any
from datetime import datetime
import hashlib
import mimetypes
from gdrive_integration import get_ai_organizer_root, get_metadata_root, ensure_safe_local_path

class ContentExtractor:
    """
    Extracts searchable content from various file formats
    Supports PDF, DOCX, TXT, Jupyter notebooks, and more
    """
    
    def __init__(self, base_dir: str = None):
        # Use Google Drive integration as primary storage root for FILE OPERATIONS
        self.base_dir = Path(base_dir) if base_dir else get_ai_organizer_root()

        # CRITICAL FIX: SQLite databases MUST be local (not on Google Drive)
        # Google Drive File Stream has known SQLite locking issues
        # Per spec: "Learning DB: ~/Documents/AI_METADATA_SYSTEM/databases/adaptive_learning.db"
        metadata_root = get_metadata_root()
        
        local_db_dir = metadata_root / "databases"
        local_db_dir.mkdir(parents=True, exist_ok=True)
        # Enforce local storage - will raise RuntimeError if unsafe
        self.db_path = ensure_safe_local_path(local_db_dir / "content_index.db")
        print(f"DEBUG: ContentExtractor DB Path: {self.db_path}")

        # Initialize extractor mapping
        self.extractors = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.pages': self._extract_pages,
            '.txt': self._extract_text,
            '.md': self._extract_markdown,
            '.ipynb': self._extract_jupyter,
            '.py': self._extract_code,
            '.js': self._extract_code,
            '.ts': self._extract_code,
            '.tsx': self._extract_code,
            '.html': self._extract_html,
            '.csv': self._extract_csv,
            '.json': self._extract_json
        }
        
        # Ensure database is initialized
        self._init_database()

    def _init_database(self):
        """Initialize content index database"""
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        
        try:
            with sqlite3.connect(self.db_path) as conn:
                conn.execute("PRAGMA journal_mode=WAL;")
                conn.execute("""
                    CREATE TABLE IF NOT EXISTS content_index (
                        file_path TEXT PRIMARY KEY,
                        file_hash TEXT,
                        content_hash TEXT,
                        extracted_text TEXT,
                        metadata TEXT,
                        extraction_method TEXT,
                        extracted_at TIMESTAMP,
                        file_size INTEGER,
                        content_length INTEGER,
                        extraction_success BOOLEAN
                    )
                """)
                print("DEBUG: Table content_index created (if not existed)")
                
                conn.execute("""
                    CREATE TABLE IF NOT EXISTS extraction_stats (
                        date DATE PRIMARY KEY,
                        files_processed INTEGER DEFAULT 0,
                        files_successful INTEGER DEFAULT 0,
                        total_content_length INTEGER DEFAULT 0,
                        avg_extraction_time REAL DEFAULT 0.0
                    )
                """)
                
                # Create full-text search index
                conn.execute("""
                    CREATE VIRTUAL TABLE IF NOT EXISTS content_fts USING fts5(
                        file_path,
                        extracted_text,
                        metadata
                    )
                """)
                conn.commit()
        except Exception as e:
            print(f"ERROR: Failed to initialize ContentExtractor database: {e}")
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Generate hash for file change detection"""
        try:
            stat = file_path.stat()
            # Use file size + modification time for quick hash
            content = f"{stat.st_size}_{stat.st_mtime}".encode()
            return hashlib.md5(content).hexdigest()
        except:
            return ""
    
    def _extract_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF files"""
        try:
            import PyPDF2
            text = ""
            metadata = {}
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Extract metadata safely
                try:
                    if reader.metadata:
                        # Force string conversion for all values to avoid IndirectObject serialization errors
                        metadata = {
                            'title': str(reader.metadata.get('/Title', '')) if reader.metadata.get('/Title') else '',
                            'author': str(reader.metadata.get('/Author', '')) if reader.metadata.get('/Author') else '',
                            'subject': str(reader.metadata.get('/Subject', '')) if reader.metadata.get('/Subject') else '',
                            'creator': str(reader.metadata.get('/Creator', '')) if reader.metadata.get('/Creator') else '',
                            'pages': len(reader.pages)
                        }
                except Exception as e:
                    # If metadata extraction fails, log it but continue with text extraction
                    print(f"Warning: Failed to extract metadata from {file_path.name}: {e}")
                    metadata = {'error': 'metadata_extraction_failed'}
                
                # Smart Page Sampling for Large PDFs
                total_pages = len(reader.pages)
                max_pages = 20
                
                pages_to_extract = range(total_pages) # Default: all
                
                if total_pages > max_pages:
                    # Construct page list: First 5, Middle 5, Last 5
                    head = range(5)
                    mid_start = total_pages // 2
                    mid = range(mid_start, mid_start + 5)
                    tail = range(total_pages - 5, total_pages)
                    
                    # Combine unique sorted indices
                    pages_to_extract = sorted(list(set(head) | set(mid) | set(tail)))
                    pages_to_extract = [p for p in pages_to_extract if 0 <= p < total_pages]
                    
                    metadata['extraction_note'] = f"Sampled {len(pages_to_extract)}/{total_pages} pages"
                
                # Extract text from selected pages
                for i in pages_to_extract:
                    try:
                        if i > 0 and i not in [p-1 for p in pages_to_extract]:
                             text += "\n\n... [Pages Skipped] ...\n\n"
                        text += reader.pages[i].extract_text() + "\n"
                    except:
                        continue
            
            return {
                'success': True,
                'text': text.strip(),
                'metadata': metadata,
                'method': 'PyPDF2' if total_pages <= max_pages else 'PyPDF2_sampled'
            }
            
        except ImportError:
            # Fallback to basic text extraction if PyPDF2 not available
            return {
                'success': False,
                'text': '',
                'metadata': {'error': 'PyPDF2 not available'},
                'method': 'none'
            }
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {'error': str(e)},
                'method': 'PyPDF2_failed'
            }
    
    def _extract_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            # Basic metadata
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            return {
                'success': True,
                'text': text.strip(),
                'metadata': metadata,
                'method': 'python-docx'
            }
            
        except ImportError:
            return {
                'success': False,
                'text': '',
                'metadata': {'error': 'python-docx not available'},
                'method': 'none'
            }
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {'error': str(e)},
                'method': 'docx_failed'
            }
    
    def _extract_pages(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Apple Pages files"""
        # Pages files are complex - this is a placeholder
        # Would need specialized extraction or conversion
        return {
            'success': False,
            'text': '',
            'metadata': {'note': 'Pages extraction requires specialized tools'},
            'method': 'pages_placeholder'
        }
    
    def _extract_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text files (with smart chunking for large files)"""
        try:
            file_size = file_path.stat().st_size
            max_size = 5 * 1024 * 1024  # 5MB threshold
            
            text = ""
            if file_size <= max_size:
                # Small file: read all
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                note = "full_content"
            else:
                # Large file: read Start, Middle, End chunks
                chunk_size = 512 * 1024  # 512KB per chunk
                
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    # 1. Start
                    text += f.read(chunk_size)
                    text += "\n\n... [Content Truncated] ...\n\n"
                    
                    # 2. Middle
                    mid_point = file_size // 2
                    f.seek(mid_point)
                    text += f.read(chunk_size)
                    text += "\n\n... [Content Truncated] ...\n\n"
                    
                    # 3. End
                    f.seek(max(0, file_size - chunk_size))
                    text += f.read()
                note = "sampled_content_large_file"

            metadata = {
                'encoding': 'utf-8',
                'lines': len(text.split('\n')),
                'extraction_note': note,
                'original_size_bytes': file_size
            }
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata,
                'method': 'direct_read' if file_size <= max_size else 'chunked_read'
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {'error': str(e)},
                'method': 'text_failed'
            }
    
    def _extract_markdown(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Markdown files"""
        result = self._extract_text(file_path)
        if result['success']:
            result['metadata']['format'] = 'markdown'
            result['method'] = 'markdown_text'
        return result
    
    def _extract_jupyter(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Jupyter notebook files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                notebook = json.load(f)
            
            text = ""
            metadata = {
                'cells': len(notebook.get('cells', [])),
                'code_cells': 0,
                'markdown_cells': 0
            }
            
            # Extract from cells
            for cell in notebook.get('cells', []):
                cell_type = cell.get('cell_type', '')
                source = cell.get('source', [])
                
                if cell_type == 'code':
                    metadata['code_cells'] += 1
                    text += "CODE CELL:\n"
                elif cell_type == 'markdown':
                    metadata['markdown_cells'] += 1
                    text += "MARKDOWN CELL:\n"
                
                # Join source lines
                if isinstance(source, list):
                    text += ''.join(source) + "\n\n"
                else:
                    text += str(source) + "\n\n"
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata,
                'method': 'jupyter_json'
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {'error': str(e)},
                'method': 'jupyter_failed'
            }
    
    def _extract_code(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from code files"""
        result = self._extract_text(file_path)
        if result['success']:
            result['metadata']['format'] = 'code'
            result['metadata']['language'] = file_path.suffix[1:]  # Remove the dot
            result['method'] = f'code_{file_path.suffix[1:]}'
        return result
    
    def _extract_html(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from HTML files"""
        try:
            from html.parser import HTMLParser
            
            class TextHTMLParser(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.text = []
                
                def handle_data(self, data):
                    self.text.append(data)
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            parser = TextHTMLParser()
            parser.feed(html_content)
            text = ' '.join(parser.text)
            
            return {
                'success': True,
                'text': text,
                'metadata': {'format': 'html'},
                'method': 'html_parser'
            }
            
        except Exception as e:
            # Fallback to direct read
            result = self._extract_text(file_path)
            result['metadata']['extraction_note'] = 'HTML parser failed, using direct read'
            return result
    
    def _extract_csv(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from CSV files"""
        try:
            import csv
            
            text_parts = []
            rows = 0
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                
                for row in reader:
                    text_parts.append(' '.join(str(cell) for cell in row))
                    rows += 1
                    
                    # Limit to first 1000 rows for performance
                    if rows >= 1000:
                        break
            
            text = '\n'.join(text_parts)
            
            return {
                'success': True,
                'text': text,
                'metadata': {'format': 'csv', 'rows_extracted': rows},
                'method': 'csv_reader'
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {'error': str(e)},
                'method': 'csv_failed'
            }
    
    def _extract_json(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to searchable text
            def flatten_json(obj, prefix=''):
                text_parts = []
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        new_prefix = f"{prefix}.{key}" if prefix else key
                        text_parts.append(f"{new_prefix}: {value}")
                        if isinstance(value, (dict, list)):
                            text_parts.extend(flatten_json(value, new_prefix))
                elif isinstance(obj, list):
                    for i, item in enumerate(obj):
                        new_prefix = f"{prefix}[{i}]" if prefix else f"[{i}]"
                        text_parts.append(f"{new_prefix}: {item}")
                        if isinstance(item, (dict, list)):
                            text_parts.extend(flatten_json(item, new_prefix))
                return text_parts
            
            text = '\n'.join(flatten_json(data))
            
            return {
                'success': True,
                'text': text,
                'metadata': {'format': 'json'},
                'method': 'json_flatten'
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'metadata': {'error': str(e)},
                'method': 'json_failed'
            }
    
    def extract_content(self, file_path: Path) -> Dict[str, Any]:
        """Extract content from a file"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                'success': False,
                'text': '',
                'metadata': {'error': 'File not found'},
                'method': 'file_not_found'
            }
        
        # Check if already extracted and up to date
        current_hash = self._get_file_hash(file_path)
        if self._is_content_cached(file_path, current_hash):
            return self._get_cached_content(file_path)
        
        # Determine extraction method
        extension = file_path.suffix.lower()
        
        if extension in self.extractors:
            result = self.extractors[extension](file_path)
        else:
            # Try to guess from MIME type
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if mime_type and mime_type.startswith('text/'):
                result = self._extract_text(file_path)
            else:
                result = {
                    'success': False,
                    'text': '',
                    'metadata': {'error': f'Unsupported file type: {extension}'},
                    'method': 'unsupported'
                }
        
        # Cache the result
        self._cache_content(file_path, current_hash, result)

        # Phase V4: Hybrid Sidecar Enrichment
        # If the file has a metadata sidecar (generated by vision/unified classifier), 
        # append its metadata to the extracted text so it gets vectorized/indexed.
        sidecar_path = file_path.with_name(f"{file_path.name}.json")
        if sidecar_path.exists():
            try:
                with open(sidecar_path, 'r', encoding='utf-8') as f:
                    sidecar_data = json.load(f)
                
                # Extract rich descriptors
                enrichment = []
                
                # Check for classification metadata
                clf = sidecar_data.get('classification', {})
                if isinstance(clf, dict):
                    if clf.get('category'):
                        enrichment.append(f"Category: {clf['category']}")
                    
                    # Add reasoning snippets
                    reasoning = clf.get('reasoning', [])
                    if isinstance(reasoning, list):
                        enrichment.extend(reasoning)
                    
                    # Add identified entities (The "Identity" part)
                    meta = clf.get('metadata', {})
                    if isinstance(meta, dict):
                        entities = meta.get('identified_entities', [])
                        if entities:
                            enrichment.append(f"Identified Entities: {', '.join(entities)}")
                        
                        desc = meta.get('description', '')
                        if desc:
                            enrichment.append(f"Description: {desc}")
                        
                        # Audio Transcript Enrichment (Phase V5)
                        transcript = meta.get('transcript') or sidecar_data.get('transcript')
                        if transcript:
                            enrichment.append(f"Audio Transcript:\n{transcript}")
                
                if enrichment:
                    sidecar_text = "\n[AI METADATA]\n" + "\n".join(enrichment)
                    result['text'] = (result['text'] + "\n" + sidecar_text).strip()
                    result['method'] = f"{result['method']}+sidecar"
                    
                    # Update cache with enriched text
                    self._cache_content(file_path, current_hash, result)
                    
            except Exception as e:
                print(f"Warning: Failed to enrich content from sidecar for {file_path.name}: {e}")
        
        return result
    
    def _is_content_cached(self, file_path: Path, file_hash: str) -> bool:
        """Check if content is already cached and up to date"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute(
                "SELECT file_hash FROM content_index WHERE file_path = ?",
                (str(file_path),)
            )
            result = cursor.fetchone()
            return result and result[0] == file_hash
    
    def _get_cached_content(self, file_path: Path) -> Dict[str, Any]:
        """Get cached content"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute("""
                SELECT extracted_text, metadata, extraction_method, extraction_success
                FROM content_index WHERE file_path = ?
            """, (str(file_path),))
            
            row = cursor.fetchone()
            if row:
                extracted_text, metadata_json, method, success = row
                metadata = json.loads(metadata_json) if metadata_json else {}
                
                return {
                    'success': bool(success),
                    'text': extracted_text or '',
                    'metadata': metadata,
                    'method': method + '_cached'
                }
        
        return {
            'success': False,
            'text': '',
            'metadata': {'error': 'Cache miss'},
            'method': 'cache_error'
        }
    
    def _cache_content(self, file_path: Path, file_hash: str, result: Dict[str, Any]):
        """Cache extracted content"""
        try:
            content_hash = hashlib.md5(result['text'].encode()).hexdigest()
            
            with sqlite3.connect(self.db_path) as conn:
                # Store in main index
                conn.execute("""
                    INSERT OR REPLACE INTO content_index 
                    (file_path, file_hash, content_hash, extracted_text, metadata, 
                     extraction_method, extracted_at, file_size, content_length, extraction_success)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    str(file_path), file_hash, content_hash, result['text'],
                    json.dumps(result['metadata']), result['method'],
                    datetime.now(), file_path.stat().st_size if file_path.exists() else 0,
                    len(result['text']), result['success']
                ))
                
                # Store in FTS index if extraction successful
                if result['success'] and result['text']:
                    conn.execute("""
                        INSERT OR REPLACE INTO content_fts (file_path, extracted_text, metadata)
                        VALUES (?, ?, ?)
                    """, (str(file_path), result['text'], json.dumps(result['metadata'])))
                
        except Exception as e:
            print(f"Error caching content for {file_path}: {e}")
    
    def search_content(self, query: str, limit: int = 10) -> List[Dict[str, Any]]:
        """Search through extracted content"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute("""
                SELECT file_path, extracted_text, metadata, 
                       rank, snippet(content_fts, 1, '<mark>', '</mark>', '...', 32) as snippet
                FROM content_fts 
                WHERE content_fts MATCH ?
                ORDER BY rank
                LIMIT ?
            """, (query, limit))
            
            results = []
            for row in cursor.fetchall():
                file_path, content, metadata_json, rank, snippet = row
                metadata = json.loads(metadata_json) if metadata_json else {}
                
                results.append({
                    'file_path': file_path,
                    'filename': Path(file_path).name,
                    'snippet': snippet,
                    'metadata': metadata,
                    'rank': rank,
                    'content_preview': content[:200] + '...' if len(content) > 200 else content
                })
            
            return results
    
    def get_extraction_stats(self) -> Dict[str, Any]:
        """Get content extraction statistics"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.execute("""
                SELECT 
                    COUNT(*) as total_files,
                    COUNT(CASE WHEN extraction_success = 1 THEN 1 END) as successful_extractions,
                    SUM(content_length) as total_content_length,
                    AVG(content_length) as avg_content_length,
                    COUNT(DISTINCT extraction_method) as extraction_methods
                FROM content_index
            """)
            
            row = cursor.fetchone()
            if row:
                total, successful, total_length, avg_length, methods = row
                return {
                    'total_files': total or 0,
                    'successful_extractions': successful or 0,
                    'success_rate': (successful / total * 100) if total > 0 else 0,
                    'total_content_length': total_length or 0,
                    'avg_content_length': round(avg_length or 0),
                    'extraction_methods': methods or 0
                }
        
        return {
            'total_files': 0,
            'successful_extractions': 0,
            'success_rate': 0,
            'total_content_length': 0,
            'avg_content_length': 0,
            'extraction_methods': 0
        }

def test_content_extractor():
    """Test the content extractor"""
    print("🧪 Testing Content Extraction Pipeline")
    print("-" * 40)
    
    # Initialize extractor
    extractor = ContentExtractor()

    # Test with sample files from AI Organizer root
    docs_path = get_ai_organizer_root()
    test_files = []
    
    # Find some files to test with
    for ext in ['.txt', '.md', '.json', '.csv', '.py', '.ipynb']:
        for file_path in docs_path.rglob(f'*{ext}'):
            if file_path.is_file():
                test_files.append(file_path)
                if len(test_files) >= 5:  # Limit to 5 test files
                    break
        if len(test_files) >= 5:
            break
    
    if not test_files:
        print("⚠️ No suitable test files found")
        return
    
    print(f"📁 Testing with {len(test_files)} files:")
    
    successful_extractions = 0
    total_content_length = 0
    
    for i, file_path in enumerate(test_files, 1):
        print(f"\n[{i}] {file_path.name}")
        
        result = extractor.extract_content(file_path)
        
        print(f"  📊 Success: {result['success']}")
        print(f"  🔧 Method: {result['method']}")
        
        if result['success']:
            successful_extractions += 1
            content_length = len(result['text'])
            total_content_length += content_length
            
            print(f"  📝 Content Length: {content_length:,} chars")
            print(f"  🔍 Preview: {result['text'][:100]}...")
            
            if result['metadata']:
                print(f"  📋 Metadata: {result['metadata']}")
        else:
            print(f"  ❌ Error: {result['metadata'].get('error', 'Unknown error')}")
    
    # Show overall stats
    stats = extractor.get_extraction_stats()
    print(f"\n📊 Extraction Statistics:")
    print(f"  Total files processed: {stats['total_files']}")
    print(f"  Successful extractions: {stats['successful_extractions']}")
    print(f"  Success rate: {stats['success_rate']:.1f}%")
    print(f"  Total content length: {stats['total_content_length']:,} chars")
    print(f"  Average content length: {stats['avg_content_length']:,} chars")
    
    # Test search functionality
    if successful_extractions > 0:
        print(f"\n🔍 Testing content search...")
        search_results = extractor.search_content("document", limit=3)
        print(f"  Found {len(search_results)} results for 'document'")
        
        for result in search_results:
            print(f"    • {result['filename']}: {result['snippet'][:100]}...")
    
    print(f"\n✅ Content extraction test completed!")

if __name__ == "__main__":
    test_content_extractor()